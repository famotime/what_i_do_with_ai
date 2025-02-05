"""
将docx文件转换为markdown文件:
1. 找到指定目录及子目录下所有docx文件；
2. 包含的图片文件保存到images目录下；
3. docx文件内容调用火山引擎模型进行文本修正后再转换为markdown文件，markdown文件内容对应位置应包含原图片链接；
4. markdown文件默认保存到对应docx文件目录下，文件名与docx文件名相同；
"""

from pathlib import Path
from docx import Document
from text_correction_with_doubao import main as correct_text

def extract_images(docx_path, image_dir):
    """从docx文件中提取图片并保存到指定目录，返回图片名称映射字典"""
    doc = Document(docx_path)
    image_dir.mkdir(parents=True, exist_ok=True)

    # 获取docx文件名（不含扩展名）作为前缀
    prefix = docx_path.stem

    # 用于存储原始文件名到新文件名的映射
    image_name_mapping = {}

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            original_name = Path(rel.target_ref).name

            # 使用docx文件名作为前缀创建新的图片文件名
            stem = Path(original_name).stem
            suffix = Path(original_name).suffix
            new_name = f"{prefix}_{stem}{suffix}"

            # 如果文件名已存在，则添加数字后缀
            if new_name in image_name_mapping.values():
                counter = 1
                while True:
                    new_name = f"{prefix}_{stem}_{counter}{suffix}"
                    if new_name not in image_name_mapping.values() and not (image_dir / new_name).exists():
                        break
                    counter += 1

            image_path = image_dir / new_name

            # 保存图片文件
            with image_path.open('wb') as f:
                f.write(image_data)

            # 记录原始文件名到新文件名的映射
            image_name_mapping[original_name] = new_name

    return image_name_mapping

def docx_to_markdown(docx_path, image_name_mapping):
    """将docx文件转换为markdown格式，使用更新后的图片链接"""
    doc = Document(docx_path)
    markdown_content = []

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1])
            markdown_content.append('#' * level + ' ' + paragraph.text)
        else:
            has_image = False
            for run in paragraph.runs:
                drawing_elem = run._r.find('./{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                if drawing_elem is not None:
                    blip_elem = drawing_elem.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if blip_elem is not None:
                        rId = blip_elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId and rId in doc.part.rels:
                            original_name = Path(doc.part.rels[rId].target_ref).name
                            # 使用映射获取新的图片名称
                            new_image_name = image_name_mapping.get(original_name, original_name)
                            markdown_content.append(f'![{new_image_name}](images/{new_image_name})')
                            has_image = True

            if paragraph.text.strip() and not has_image:
                markdown_content.append(paragraph.text)

    return '\n\n'.join(markdown_content)

def process_files(input_dir, endpoint_id, api_host, system_message, file_type='docx', skip_existing=False):
    """
    处理指定目录下的docx或md文件

    Args:
        input_dir (str): 输入目录路径
        endpoint_id (str): 火山引擎模型端点ID
        api_host (str): API主机地址
        system_message (str): 系统提示信息
        file_type (str): 文件类型，'docx'或'md'
        skip_existing (bool): 是否跳过已存在的目标文件
    """
    input_dir = Path(input_dir)

    # 根据文件类型选择处理方式
    if file_type.lower() == 'docx':
        # 获取所有docx文件列表，排除临时文件(~$开头)和包含"_combined"的文件
        docx_files = [f for f in input_dir.rglob('*.docx')
                     if '_combined' not in f.stem and not f.stem.startswith('~$')]
        total_files = len(docx_files)
        print(f"\n共找到 {total_files} 个docx文件待处理")

        processed_count = 0
        skipped_count = 0

        for index, docx_file in enumerate(docx_files, 1):
            output_path = docx_file.with_suffix('.md')

            # 检查目标文件是否已存在
            if skip_existing and output_path.exists():
                print(f"\n[{index}/{total_files}] 跳过已存在文件: {output_path}")
                skipped_count += 1
                continue

            print(f"\n[{index}/{total_files}] 处理文件: {docx_file}")
            image_dir = output_path.parent / 'images'

            try:
                # 提取图片并获取文件名映射
                image_name_mapping = extract_images(docx_file, image_dir)
                # 使用文件名映射生成markdown内容
                markdown_content = docx_to_markdown(docx_file, image_name_mapping)

                print("- 优化文本内容...")
                optimized_content = correct_text(system_message, endpoint_id, api_host, markdown_content, max_length=1500)

                output_path.write_text(optimized_content, encoding='utf-8')
                print(f"✓ 完成 ({index}/{total_files}) - 已保存到: {output_path}")
                processed_count += 1
            except Exception as e:
                print(f"× 处理失败: {str(e)}")

    elif file_type.lower() == 'md':
        # 获取所有md文件列表，排除包含"_combined"的文件
        md_files = [f for f in input_dir.rglob('*.md') if '_combined' not in f.stem]
        total_files = len(md_files)
        print(f"\n共找到 {total_files} 个md文件待处理")

        processed_count = 0
        skipped_count = 0

        for index, md_file in enumerate(md_files, 1):
            backup_file = md_file.with_suffix('.md.bak')

            # 检查备份文件是否已存在
            if skip_existing and backup_file.exists():
                print(f"\n[{index}/{total_files}] 跳过已处理文件: {md_file}")
                skipped_count += 1
                continue

            print(f"\n[{index}/{total_files}] 处理文件: {md_file}")

            try:
                markdown_content = md_file.read_text(encoding='utf-8')

                print("- 优化文本内容...")
                optimized_content = correct_text(system_message, endpoint_id, api_host, markdown_content, max_length=1500)

                md_file.rename(backup_file)

                # 保存优化后的内容到原文件
                md_file.write_text(optimized_content, encoding='utf-8')
                print(f"✓ 完成 ({index}/{total_files}) - 原文件已备份为: {backup_file}，优化文件保存到: {md_file}")
                processed_count += 1
            except Exception as e:
                print(f"× 处理失败: {str(e)}")
    else:
        raise ValueError(f"不支持的文件类型: {file_type}，请使用 'docx' 或 'md'")

    print(f"\n处理完成！")
    print(f"总文件数: {total_files}，其中成功处理: {processed_count}，已跳过: {skipped_count}")

if __name__ == "__main__":
    # 火山引擎模型配置
    endpoint_id = "ep-20241201202141-xghlt"  # doubao-pro-4k
    api_host = "ark.cn-beijing.volces.com"   # 华北 2 (北京) 服务器

    # 设置文本优化的系统提示
    system_message = """
    你是一名精通中文的语言专家，请对文本进行语法修正，适当修改词句，按照文章含义合理拆分段落，
    让整体文章内容更流畅，词句更偏向书面写作风格，但所有修改要求贴合原意，不要做大的改动。
    请保持markdown格式的图片链接不变。
    """

    # 设置输入目录和文件类型
    input_dir = r"D:\小汤汁茶馆知识星球哈\精华内容2019.8-2024.10（花了我兼职88元！呜呜）\1  精华中的精华哈"
    file_type = "docx"  # 或 "md"
    skip_existing = True  # 是否跳过已存在的文件

    # 处理文件
    process_files(input_dir, endpoint_id, api_host, system_message, file_type, skip_existing)
